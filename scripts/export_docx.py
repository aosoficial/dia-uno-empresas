#!/usr/bin/env python3
"""Exporta la documentación combinada a formato DOCX.

Lee build/outputs/master_combined.md (generado por build_docs.py) y lo
convierte a DOCX. Intenta usar python-docx; si no está instalado, genera
un DOCX mínimo usando zipfile y la estructura Open XML básica.

Uso:
    python scripts/build_docs.py   # primero generar el .md combinado
    python scripts/export_docx.py

Salida:
    build/outputs/master.docx
"""
from __future__ import annotations

import re
import sys
import zipfile
from io import BytesIO
from pathlib import Path
from xml.etree.ElementTree import Element, SubElement, tostring

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "build" / "outputs"
INPUT_FILE = OUTPUT_DIR / "master_combined.md"
OUTPUT_FILE = OUTPUT_DIR / "master.docx"

# --- Namespace para Open XML ---
WPML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"


# ─────────────────────────────────────────────
# Funciones auxiliares para parseo Markdown
# ─────────────────────────────────────────────

def _parse_md_blocks(text: str) -> list[dict]:
    """Parsea texto Markdown en bloques simples.

    Cada bloque es un dict con 'type' y 'content'.
    Tipos: heading, paragraph, list_item, code_block, separator.
    """
    blocks: list[dict] = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Separador
        if re.match(r"^---+\s*$", line):
            blocks.append({"type": "separator", "content": ""})
            i += 1
            continue

        # Bloque de código
        if line.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # saltar cierre ```
            blocks.append({"type": "code_block", "content": "\n".join(code_lines)})
            continue

        # Encabezado
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            blocks.append({"type": "heading", "level": level, "content": heading_match.group(2).strip()})
            i += 1
            continue

        # Elemento de lista
        list_match = re.match(r"^(\s*[-*+]|\s*\d+\.)\s+(.+)$", line)
        if list_match:
            blocks.append({"type": "list_item", "content": list_match.group(2).strip()})
            i += 1
            continue

        # Línea vacía — saltar
        if not line.strip():
            i += 1
            continue

        # Párrafo: acumular líneas consecutivas no vacías
        para_lines: list[str] = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            if not next_line.strip():
                break
            if re.match(r"^#{1,6}\s+", next_line):
                break
            if re.match(r"^---+\s*$", next_line):
                break
            if next_line.startswith("```"):
                break
            if re.match(r"^(\s*[-*+]|\s*\d+\.)\s+", next_line):
                break
            para_lines.append(next_line)
            i += 1

        blocks.append({"type": "paragraph", "content": " ".join(para_lines)})

    return blocks


def _strip_md_inline(text: str) -> str:
    """Elimina marcado inline básico (bold, italic, code) para texto plano."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


# ─────────────────────────────────────────────
# Ruta 1: python-docx disponible
# ─────────────────────────────────────────────

def _export_with_python_docx(text: str) -> None:
    """Genera DOCX usando la librería python-docx."""
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore

    doc = Document()
    blocks = _parse_md_blocks(text)

    for block in blocks:
        btype = block["type"]
        content = block.get("content", "")

        if btype == "heading":
            level = min(block.get("level", 1), 4)  # python-docx soporta heading 1-9
            doc.add_heading(_strip_md_inline(content), level=level)

        elif btype == "paragraph":
            doc.add_paragraph(_strip_md_inline(content))

        elif btype == "list_item":
            doc.add_paragraph(_strip_md_inline(content), style="List Bullet")

        elif btype == "code_block":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        elif btype == "separator":
            # Línea visual como párrafo delgado
            doc.add_paragraph("─" * 60)

    doc.save(str(OUTPUT_FILE))


# ─────────────────────────────────────────────
# Ruta 2: fallback sin python-docx (Open XML manual)
# ─────────────────────────────────────────────

def _make_xml_paragraph(text: str, *, bold: bool = False, font_size: int | None = None,
                        style: str | None = None) -> Element:
    """Crea un elemento <w:p> con un run de texto."""
    p = Element(f"{{{WPML}}}p")
    if style:
        pPr = SubElement(p, f"{{{WPML}}}pPr")
        pStyle = SubElement(pPr, f"{{{WPML}}}pStyle")
        pStyle.set(f"{{{WPML}}}val", style)
    r = SubElement(p, f"{{{WPML}}}r")
    if bold or font_size:
        rPr = SubElement(r, f"{{{WPML}}}rPr")
        if bold:
            SubElement(rPr, f"{{{WPML}}}b")
        if font_size:
            sz = SubElement(rPr, f"{{{WPML}}}sz")
            sz.set(f"{{{WPML}}}val", str(font_size))
    t = SubElement(r, f"{{{WPML}}}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return p


def _export_fallback(text: str) -> None:
    """Genera un DOCX mínimo válido usando zipfile + Open XML."""
    blocks = _parse_md_blocks(text)

    # Construir document.xml
    body_el = Element(f"{{{WPML}}}body")
    document_el = Element(f"{{{WPML}}}document")
    document_el.append(body_el)

    heading_sizes = {1: 48, 2: 36, 3: 28, 4: 24, 5: 22, 6: 20}

    for block in blocks:
        btype = block["type"]
        content = block.get("content", "")

        if btype == "heading":
            level = block.get("level", 1)
            size = heading_sizes.get(level, 20)
            body_el.append(_make_xml_paragraph(
                _strip_md_inline(content), bold=True, font_size=size))

        elif btype == "paragraph":
            body_el.append(_make_xml_paragraph(_strip_md_inline(content)))

        elif btype == "list_item":
            body_el.append(_make_xml_paragraph(
                f"  \u2022  {_strip_md_inline(content)}"))

        elif btype == "code_block":
            for code_line in content.split("\n"):
                body_el.append(_make_xml_paragraph(code_line, font_size=18))

        elif btype == "separator":
            body_el.append(_make_xml_paragraph("\u2500" * 60))

    doc_xml = tostring(document_el, encoding="unicode")
    # Añadir declaración XML
    doc_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + doc_xml

    # content_types.xml
    types_el = Element("Types")
    types_el.set("xmlns", CT_NS)
    default_rels = SubElement(types_el, "Default")
    default_rels.set("Extension", "rels")
    default_rels.set("ContentType", "application/vnd.openxmlformats-package.relationships+xml")
    default_xml = SubElement(types_el, "Default")
    default_xml.set("Extension", "xml")
    default_xml.set("ContentType", "application/xml")
    override = SubElement(types_el, "Override")
    override.set("PartName", "/word/document.xml")
    override.set("ContentType",
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml")
    ct_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + tostring(types_el, encoding="unicode")

    # _rels/.rels
    rels_el = Element("Relationships")
    rels_el.set("xmlns", REL_NS)
    rel = SubElement(rels_el, "Relationship")
    rel.set("Id", "rId1")
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument")
    rel.set("Target", "word/document.xml")
    rels_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + tostring(rels_el, encoding="unicode")

    # word/_rels/document.xml.rels (vacío pero necesario)
    doc_rels_el = Element("Relationships")
    doc_rels_el.set("xmlns", REL_NS)
    doc_rels_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + tostring(doc_rels_el, encoding="unicode")

    # Escribir ZIP
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", ct_xml)
        zf.writestr("_rels/.rels", rels_xml)
        zf.writestr("word/document.xml", doc_xml)
        zf.writestr("word/_rels/document.xml.rels", doc_rels_xml)

    OUTPUT_FILE.write_bytes(buf.getvalue())


# ─────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────

def main() -> int:
    """Punto de entrada principal."""
    if not INPUT_FILE.exists():
        print(f"No se encontró {INPUT_FILE.relative_to(ROOT)}")
        print("Ejecuta primero: python scripts/build_docs.py")
        return 1

    print(f"Leyendo: {INPUT_FILE.relative_to(ROOT)}")
    md_text = INPUT_FILE.read_text(encoding="utf-8")
    print(f"  {len(md_text)} caracteres leídos.")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Intentar python-docx; si no está, usar fallback
    try:
        import docx  # type: ignore  # noqa: F401
        print("Usando python-docx para la conversión...")
        _export_with_python_docx(md_text)
        method = "python-docx"
    except ImportError:
        print("python-docx no disponible. Usando generador Open XML integrado...")
        _export_fallback(md_text)
        method = "Open XML fallback"

    size_kb = OUTPUT_FILE.stat().st_size / 1024
    print(f"\nDocumento generado: {OUTPUT_FILE.relative_to(ROOT)}")
    print(f"  Tamaño: {size_kb:.1f} KB")
    print(f"  Método: {method}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
